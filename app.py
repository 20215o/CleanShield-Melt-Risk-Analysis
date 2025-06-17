import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_squared_error, r2_score
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Set page configuration with wide layout and custom title
st.set_page_config(layout="wide", page_title="CleanShield Albedo & Melt Risk Analysis")

# Custom CSS for professional styling
st.markdown(
    """
    <style>
    .sidebar .sidebar-content {
        background: #1a2a44;
        color: #e0e7ff;
    }
    .main {
        background: linear-gradient(to bottom, #2e4057, #1a2a44);
        color: #e0e7ff;
        padding: 20px;
    }
    .stApp {
        background-color: transparent;
    }
    h1 {
        color: #00cc99;
        font-size: 2.5em;
        font-weight: bold;
        text-shadow: 2px 2px 4px #000000;
    }
    h2, h3 {
        color: #1e90ff;
        font-weight: bold;
    }
    .stButton>button {
        background-color: #00cc99;
        color: white;
        border: none;
        padding: 10px 20px;
        border-radius: 5px;
        font-weight: bold;
    }
    .stButton>button:hover {
        background-color: #009970;
    }
    .stText {
        color: #e0e7ff;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Detailed Problem Statement
st.title("CleanShield Albedo & Melt Risk Analysis")
st.markdown("""
### The Problem
Glaciers worldwide are melting at an alarming rate due to rising global temperatures, air pollution (notably PM2.5 particulates), and reduced albedo caused by soot and dust deposition. This threatens freshwater supplies for billions, disrupts ecosystems, and contributes to sea level rise, endangering coastal communities. By 2050, glacier loss could reduce water availability by up to 20% in regions like the Himalayas and Andes, worsening droughts and food insecurity. CleanShield predicts melt risk using environmental data and provides actionable community solutions.

### How to Use
1. **Data Source**: Select curated data, upload a CSV, or provide a URL with columns: pm25, temperature, humidity, elevation, albedo, solar_radiation, wind_speed, snow_depth, precipitation, melt_risk.
2. **Explore**: Preview data and train the model.
3. **Visualize**: Analyze trends with multiple plots.
4. **Predict & Act**: Obtain risk predictions, alerts, and a detailed action plan.

**Note**: Melt risk thresholds: >75 (High), >50 (Medium), ≤50 (Low).
""", unsafe_allow_html=True)

# Sidebar for settings
st.sidebar.header("Configuration Panel")
upload_option = st.sidebar.radio("Select Data Source", ["Use Curated Data", "Upload CSV", "URL"])

# Initialize df as None
df = None

# Data loading outside cached function
if upload_option == "Upload CSV":
    uploaded_file = st.sidebar.file_uploader("Upload CSV File", type=["csv"], help="Upload a CSV with columns: pm25, temperature, humidity, elevation, albedo, solar_radiation, wind_speed, snow_depth, precipitation, melt_risk.")
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)
elif upload_option == "URL":
    url = st.sidebar.text_input("Enter CSV URL", help="Provide a public URL to a CSV file (e.g., Google Drive link).")
    if url:
        df = pd.read_csv(url)
else:  # Use Curated Data
    try:
        df = pd.read_csv("curated_glacier_observations.csv")
    except FileNotFoundError:
        st.sidebar.error("Curated data file 'curated_glacier_observations.csv' not found. Please upload a CSV or provide a URL.")
        df = None

# Cache data processing separately
@st.cache_data
def process_data(_df):
    if _df is not None and not all(col in _df.columns for col in ['pm25', 'temperature', 'humidity', 'elevation', 'albedo', 'solar_radiation', 'wind_speed', 'snow_depth', 'precipitation', 'melt_risk']):
        st.error("Missing required columns. Ensure CSV includes: pm25, temperature, humidity, elevation, albedo, solar_radiation, wind_speed, snow_depth, precipitation, melt_risk.")
        return None
    return _df

if df is not None:
    df = process_data(df)
else:
    st.warning("No data loaded. Please upload a CSV, provide a URL, or ensure 'curated_glacier_observations.csv' exists.")

# Define features at a higher scope
features = ['pm25', 'temperature', 'humidity', 'elevation', 'albedo', 
            'solar_radiation', 'wind_speed', 'snow_depth', 'precipitation']

if df is not None:
    st.subheader("Data Overview")
    st.markdown("<h3>Data Preview</h3>", unsafe_allow_html=True)
    st.write("A snapshot of environmental factors and melt risk data.", df.head())
    st.caption("Columns: pm25 (µg/m³), temperature (°C), humidity (%), elevation (m), albedo, solar_radiation (W/m²), wind_speed (m/s), snow_depth (cm), precipitation (mm), melt_risk (index).")

    # Model training and evaluation
    st.subheader("Model Development")
    st.markdown("<h3>Model Training</h3>", unsafe_allow_html=True)
    st.write("Train a Random Forest model to predict melt risk based on environmental factors.")
    if st.button("Train Model"):
        X = df[features]
        y = df['melt_risk']
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

        model = RandomForestRegressor(n_estimators=250, max_depth=15, min_samples_split=5, random_state=42)
        model.fit(X_train, y_train)
        
        y_pred = model.predict(X_test)
        mse = mean_squared_error(y_test, y_pred)
        r2 = r2_score(y_test, y_pred)
        
        st.session_state.model = model  # Store model in session state
        st.success("Model trained successfully!")  # Confirm training
        st.write(f"**Mean Squared Error**: {mse:.2f}")
        st.write(f"**R² Score**: {r2:.2f}")

    # Visualizations
    st.subheader("Data Visualization")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("<h3>Correlation Heatmap</h3>", unsafe_allow_html=True)
        st.write("Visualizes the correlation between environmental factors and melt risk.")
        plt.figure(figsize=(8, 6))
        sns.heatmap(df.corr(), annot=True, cmap='coolwarm', vmin=-1, vmax=1)
        st.pyplot(plt)

    with col2:
        st.markdown("<h3>Melt Risk Distribution</h3>", unsafe_allow_html=True)
        st.write("Illustrates the distribution of melt risk values.")
        plt.figure(figsize=(8, 6))
        sns.histplot(df['melt_risk'], bins=30, kde=True)
        st.pyplot(plt)

    col3, col4 = st.columns(2)
    with col3:
        st.markdown("<h3>Temperature vs. Melt Risk</h3>", unsafe_allow_html=True)
        st.write("Displays the relationship between temperature and melt risk.")
        plt.figure(figsize=(8, 6))
        plt.scatter(df['temperature'], df['melt_risk'], alpha=0.5)
        plt.xlabel("Temperature (°C)")
        plt.ylabel("Melt Risk")
        st.pyplot(plt)

    with col4:
        st.markdown("<h3>Box Plot of Key Variables</h3>", unsafe_allow_html=True)
        st.write("Highlights outliers in albedo and PM2.5 levels.")
        plt.figure(figsize=(8, 6))
        sns.boxplot(data=df[['albedo', 'pm25']])
        plt.ylabel("Value")
        st.pyplot(plt)

    # Additional Visualizations
    st.subheader("Advanced Visualizations")
    if 'date' in df.columns:
        df['date'] = pd.to_datetime(df['date'])
        st.markdown("<h3>Melt Risk Over Time</h3>", unsafe_allow_html=True)
        st.write("Tracks melt risk trends over time.")
        plt.figure(figsize=(10, 6))
        plt.plot(df['date'], df['melt_risk'], label="Melt Risk", color='#1e90ff')
        plt.xlabel("Date")
        plt.ylabel("Melt Risk")
        plt.title("Melt Risk Trend Over Time")
        plt.legend()
        st.pyplot(plt)
    else:
        st.markdown("<h3>Simulated Melt Risk Trend</h3>", unsafe_allow_html=True)
        st.write("Simulated trend of melt risk over time using index.")
        plt.figure(figsize=(10, 6))
        plt.plot(range(len(df)), df['melt_risk'], label="Melt Risk", color='#1e90ff')
        plt.xlabel("Time (Simulated Index)")
        plt.ylabel("Melt Risk")
        plt.title("Simulated Melt Risk Trend")
        plt.legend()
        st.pyplot(plt)

    st.markdown("<h3>Pair Plot</h3>", unsafe_allow_html=True)
    st.write("Explores relationships between key variables (pm25, temperature, albedo, melt_risk).")
    plt.figure(figsize=(10, 8))
    sns.pairplot(df[['pm25', 'temperature', 'albedo', 'melt_risk']])
    st.pyplot(plt)

    st.markdown("<h3>Violin Plot</h3>", unsafe_allow_html=True)
    st.write("Shows distribution and density of melt risk by temperature quartiles.")
    plt.figure(figsize=(10, 6))
    sns.violinplot(x=pd.qcut(df['temperature'], 4, labels=["Low", "Medium", "High", "Very High"]), y=df['melt_risk'])
    plt.xlabel("Temperature Quartiles")
    plt.ylabel("Melt Risk")
    plt.title("Melt Risk Distribution by Temperature")
    st.pyplot(plt)

    # Prediction and Solution System
    st.subheader("Risk Assessment & Mitigation")
    if 'model' in st.session_state:
        st.markdown("<h3>Melt Risk Prediction</h3>", unsafe_allow_html=True)
        st.write("Adjust the sliders below to predict melt risk and receive tailored solutions.")
        input_data = {}
        for feature in features:
            input_data[feature] = st.slider(f"{feature.capitalize()}", float(df[feature].min()), float(df[feature].max()), float(df[feature].mean()), help=f"Adjust the {feature} value.")
        input_df = pd.DataFrame([input_data])
        prediction = st.session_state.model.predict(input_df)[0]
        st.write(f"**Predicted Melt Risk**: {prediction:.2f}")

        # Alert and Solution System
        if prediction > 75:
            st.error("**High Melt Risk Alert (>75)**: Immediate action is critical!")
            st.write("**Recommended Actions**: Install shading structures, enforce PM2.5 controls, deploy artificial snowmaking.")
        elif prediction > 50:
            st.warning("**Medium Melt Risk Alert (>50)**: Preventive measures are advised.")
            st.write("**Recommended Actions**: Monitor temperature and albedo, reduce pollution, plan water management.")
        else:
            st.success("**Low Melt Risk Alert (≤50)**: Maintain ongoing vigilance.")
            st.write("**Recommended Actions**: Continue monitoring, promote sustainability, educate communities.")

        # Comparative Analysis
        avg_melt_risk = df['melt_risk'].mean()
        st.write(f"**Comparison**: Predicted ({prediction:.2f}) vs. Average ({avg_melt_risk:.2f}).")
        if prediction > avg_melt_risk:
            st.write("**Insight**: Above average—prioritize mitigation strategies.")
        else:
            st.write("**Insight**: Below average—focus on preventive measures.")

        # Real-Time Alert Simulation
        if prediction > 75:
            st.error("**Urgent Notification**: High melt risk detected! Take action immediately.")

        # Enhanced Action Plan
        if st.button("Generate Action Plan"):
            st.write("Generating action plan...")  # Debug message
            action_plan = f"""
            **Community Action Plan for Glacier Melt Mitigation**
            **Risk Assessment**
            - Predicted Melt Risk: {prediction:.2f} ({'High (>75)' if prediction > 75 else 'Medium (>50)' if prediction > 50 else 'Low (≤50)'})
            - This assessment is derived from current environmental data, highlighting the urgency of mitigation efforts.
            **Recommended Actions**
            - **{('Install shading structures to reduce solar absorption by glaciers.' if prediction > 75 else 'Monitor temperature and albedo regularly to detect changes early.' if prediction > 50 else 'Continue regular monitoring to maintain baseline data.')}**
              - *Purpose*: Reduces heat absorption or tracks trends for timely intervention.
              - *Method*: Use lightweight, reflective materials for shading; deploy real-time albedo sensors.
              - *Impact*: Can reduce melt by 10-20% in high-risk areas; enables early detection.
            - **{('Enforce air pollution controls to limit PM2.5 deposition on glaciers.' if prediction > 75 else 'Reduce local pollution sources to improve air quality.' if prediction > 50 else 'Promote sustainable practices to minimize environmental impact.')}**
              - *Purpose*: Prevents albedo reduction or enhances air quality.
              - *Method*: Implement strict emission regulations; adopt clean energy solutions.
              - *Impact*: May increase albedo by 5-15%, lowering melt risk over time.
            - **{('Deploy artificial snowmaking to replenish glacier mass.' if prediction > 75 else 'Plan water management strategies to adapt to melt changes.' if prediction > 50 else 'Educate communities on sustainable water use.')}**
              - *Purpose*: Restores reflective surfaces or mitigates flood risks.
              - *Method*: Use treated water with snow cannons; develop flood defenses or conduct workshops.
              - *Impact*: Restores 5-10% of mass annually; reduces flooding by up to 30%.
            **Local Resources**
            - Collaborate with environmental organizations for expertise and funding.
            - Seek government grants to support large-scale projects.
            - Engage community leaders to ensure local support and participation.
            """
            st.markdown(action_plan, unsafe_allow_html=True)
            # Word document
            doc = Document()
            doc.add_heading("Community Action Plan for Glacier Melt Mitigation", 0)
            doc.add_paragraph(action_plan.replace("<b>", "").replace("</b>", "").replace("*", ""))
            docx_file = "action_plan.docx"
            doc.save(docx_file)
            with open(docx_file, "rb") as f:
                st.download_button(label="Download Action Plan (Word)", data=f, file_name=f"action_plan_{prediction:.2f}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # PDF document
            pdf_file = "action_plan.pdf"
            doc = SimpleDocTemplate(pdf_file, pagesize=letter)
            styles = getSampleStyleSheet()
            custom_style = ParagraphStyle(
                name='CustomStyle',
                parent=styles['Normal'],
                fontSize=12,
                leading=14,
                textColor=colors.black
            )
            # Simplify text for PDF by removing HTML tags and using plain text with styles
            plain_text = action_plan.replace("<b>", "").replace("</b>", "").replace("*", "").replace("\n", " ")
            story = [
                Paragraph("Community Action Plan for Glacier Melt Mitigation", styles['Heading1']),
                Spacer(1, 12),
                Paragraph(plain_text, custom_style)
            ]
            doc.build(story)
            with open(pdf_file, "rb") as f:
                st.download_button(label="Download Action Plan (PDF)", data=f, file_name=f"action_plan_{prediction:.2f}.pdf", mime="application/pdf")

            st.write("Download the action plan in Word or PDF format to share with your community.")
            st.stop()  # Prevent page reload
    else:
        st.write("Model not trained. Please click 'Train Model' to proceed.")